#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Nov  7 18:21:10 2018
@author: glenabastillas

WRITER SCRIPT
-------------------------------------------------------------------------------
Glenn Abastillas | November 7, 2018 | writer.py
-------------------------------------------------------------------------------

Description:

These methods are used to build a document report. The Report Writer package
follows the paradigm inserting variable data into a pre-drafted, pre-formatted
Word document. The pre-drafted, pre-formatted Word document (document) should
contain variables in the format `$ABC`.

These variables to be replaced are to be defined in a separate configuration
spreadsheet that contains, at minimum, the following columns:

    variable name   <- Variable name in document $VAR
    dataset         <- d90 (90-day period), dn (current), dp (previous)
    method          <- Function used to derive $VAR value
    table           <- Table function to use (see calculate package)
    column          <- Name of column in data
    type            <- Data type to output (e.g., numeric, string)
    value           <- Initially blank. To be filled out with infuse()
    kwargs          <- Keyword arguments in the format {'key': 'value'}

Dependencies:
    pandas
    docx
    calculate (custom package)
    database_access (custom package)
"""

from database import BoardDB, config
from calculate import date_week, date_fortnight, date_midmonth, date_month
from calculate import ccr_detect_topics, ccr_augment, date_90
from calculate import ccr_top_compliments_concerns
from calculate import date_format, date_delta

from docx import Document
from docx.shared import Inches

import pandas as pd
import calculate
import json
import re
import pickle

try:
    from tqdm import tqdm
except ImportError:
    pass

database = BoardDB()
document = Document(config.paths.document)
variables = pd.read_excel(config.paths.variables)
this = config.script.writer


def __get_date_func(delta):
    """
    Return a function that determines start and end dates for a `delta`.

    Parameters
    ----------
        delta (str): time period of analysis (e.g., week, two weeks).

    Returns
    -------
        Method that can determine a start and end date.

    Example
    -------
        >>> __get_date_func("week")
        >>> # Returns the date_week() function
    """
    periods = [re.compile(_, re.I) for _ in this.get_date_func.patterns]
    one_week, two_weeks, mid_month, one_month = periods

    delta = str(delta)

    if mid_month.findall(delta):
        return date_midmonth
    else:
        if two_weeks.findall(delta):
            return date_fortnight
        elif one_month.findall(delta):
            return date_month
        elif one_week.findall(delta):
            return date_week
        else:
            message = "Only periods of `week` or `month` are allowed."
            raise AttributeError(message)


def _subset_by_date(df, start, end, column='ResponseDateTime'):
    """ Filter the data to meet the date conditions. """
    T = pd.Timestamp
    X = df[column].apply(T)
    A = X >= date_format(start, asdatetime=True)
    B = X < date_format(end, asdatetime=True)
    return df[A & B].copy()


def query(delta='week', end=config.today, **kwargs):
    """
    Return a dict of DataFrames derived from a SQL query and an analysis.
    These DataFrames are used to supply a template document with values
    for the purposes of automating document creation.

    Parameters
    ----------
        delta (str): time period to compare (e.g., week, two weeks)
        end (datetime) End date. Default config.today

        kwargs
        ------

        func (function): Query function from database_access
        fn (str): fn (str): File name for query results
        save (boolean): Save query results?
        method (str): Type of A11 CX Domain calculation to use

    Returns
    -------
        dict of DataFrames for each time period of interest and each
        type of data involved including data from:

                DataFrame Elements
                ------------------
                    (d90) 90-Day period by ResponseDateTime
                    (dn)  Current period of comparison by ResponseDateTime
                    (dp)  Previous period of comparison by ResponseDateTime
                    (aug) Augmented DataFrame for Topics
                    (a11) A-11 CX Domain data from VEO DB
                    (tot) SurveyType data for Trust over Time time series
                    (mp0) Current month's A11 scores
                    (mp1) Previous 1 month's A11 scores
                    (mp2) Previous 2 month's A11 scores
                    (mp3) Previous 3 month's A11 scores
                    (mp4) Previous 4 month's A11 scores

                Non-DataFrame Elements
                ----------------------

                    Dictionaries
                    ------------
                        (ccr) Compliments, Concerns, Recommendation

                    Strings
                    -------
                        (s90) Start date for 90-day period (not a DataFrame)
                        (e90) End date for 90-day period (not a DataFrame)
                        (sn) Start date for current period (not a DataFrame)
                        (en) End date for current period (not a DataFrame)
                        (sp) Start date for previous period (not a DataFrame)
                        (ep) End date for previous period (not a DataFrame)

    Notes
    -----
        A11 CX Domain calculations are either by Question "Q"/"q" or by
        Response "R"/"r". Default is by question.

    Example
    -------
        >>> # Using default parameters
        >>> data = query()
        >>>
        >>> # Specifying a time period (delta)
        >>> data = query('month')
        >>>
        >>> # Specifying a different end date (Y, m, d)
        >>> data = query('month', (2019, 1, 1))
        >>>
    """
    # Get Start and End dates for each time period
    end = date_format(end, asdatetime=True)
    date_column = kwargs.pop("date_column", "ResponseDateTime")

    # Base data uses a max 155 day time span (5 months * 31 days)
    pull = kwargs.pop("func", database.pull)
    d120 = pull(end=date_format(end), diff=kwargs.pop("diff", 155))
    d120[date_column] = d120[date_column].astype("datetime64[ns]")
    d120['Trust'] = d120['Trust'].astype('int32')

    if d120.shape[0] == 0:
        error_message = this.query.error
        raise Exception(error_message)

    date_func = __get_date_func(delta)

    end90, start90 = date_90(end, astuple=False)
    endn, startn, endp, startp = date_func(end, asdatetime=False)

    # Create the 90-day, current, and previous period datasets
    dates = [(start90, end90), (startn, endn), (startp, endp)]
    d90, dn, dp = [_subset_by_date(d120, start, end) for start, end in dates]

    # Create the augmented (aug) and feedback (Compliments, Concerns) datasets
    comments = ccr_augment(dn)
    aug = ccr_detect_topics(comments)
    compliments, concerns = ccr_top_compliments_concerns(aug)

    # Create the special (a11, tot) datasets
    method = kwargs.pop("method", "r")
    a11 = calculate.calculate_a11(d90, method=method)

#    tot = database.pull_tot()  # *calculate.date_ytd())
    tot = calculate.normalize_survey_type(database.pull_tot())
    tot = calculate.calculate_tot(tot)

    ccr = dict(Compliment=compliments, Concern=concerns)

    # These datasets are for the previous 3 months (each) for Trust score
    number_of_months = kwargs.pop("number_of_months", 4)
    scores_by_month = []
    names_for_month = []
    end_, start_ = endn, startn

    while number_of_months > 0:
        A = database.pull_a11(start=start_, end=end_)
        scores = calculate.calculate_a11(A, method=method)
                
        scores_by_month.append(scores)
        names_for_month.append(end_)
        
        end_, start_ = date_delta(start_, months=1, adjust=False)
        number_of_months -= 1

    # Current month to the last 4 months
    mp0, mp1, mp2, mp3 = scores_by_month
    np0, np1, np2, np3 = names_for_month

    data = dict(d90=d90, dn=dn, dp=dp, aug=aug, ccr=ccr,
                a11=a11, tot=tot, s90=start90, e90=end90,
                sn=startn, en=endn, sp=startp, ep=endp,
                mp0=mp0, mp1=mp1, mp2=mp2, mp3=mp3,
                np0=np0, np1=np1, np2=np2, np3=np3,)
    
    save = kwargs.pop("save", False)

    fn = kwargs.pop('fn', None)
    if fn is not None:
        save = True

    if save:
        creation_date, end = [f"{_:%Y%m%d}" for _ in [config.today, end]]
        fn_dict = dict(creation_date=creation_date, fn=fn, file_end_date=end)

        if fn is None:
            fn = config.filenames.board_data.format(**fn_dict)
        else:
            fn = config.filenames.board_user.format(**fn_dict)

        with open(config.paths.data / fn, 'wb') as data_out:
            pickle.dump(data, data_out, protocol=3)

    return data


def analyze(variables, data, a11='q', package=calculate, progress_bar=False):
    """
    Calculate values for each `$VARIABLE` in variables spreadsheet and up-
    date the variables DataFrame with values.

    Parameters
    ----------
        variables (DataFrame): Variables, methods, and blank values.
        data (dict): Query results from `query`.
        package (py): Package with functions. Default is `calculate`.
        progress_bar (bool): Show progress bar for this method.
        
    Notes
    -----
        Requires the use of a `variables` DataFrame defined outside of this
        script and a `calculate` or equivalent package. Custom packages can be
        substituted in for the default `calculate` package as long as all of 
        the methods defined in `variables` also exist in the `calculate` or 
        custom package.
    """
    iterable = variables.iterrows()
    if 'tqdm' in locals() and progress_bar:
        iterable = tqdm(variables.iterrows())
    elif progress_bar:
        warning = ("Warning: `tqdm` is not an available package for import. "
                   "Will continue analysis without progress bar.")        
        print(warning)    
    
    for index, row in iterable:

        # If this row is deactivated, check the next row
        if not row.active:
            continue

#        print(row.variable, row.dataset, row.method)

        # Check for presence of method in package
        if pd.notnull(row.method) and hasattr(package, row.method):

            function = getattr(package, row.method)
            kwargs = json.loads(row.kwargs)

            # Run methods on current (dn) and previous (dp) datasets
            if ('dn' in row.dataset) and ('dp' in row.dataset):
                current, before = data['dn'], data['dp']

                if pd.notnull(row.table):
                    current = getattr(package, row.table)(current, row.param1)
                    before = getattr(package, row.table)(before, row.param1)

                value = function(current, before, **kwargs)

            # Run methods on single dataset (e.g., d90, aug, std)
            elif pd.notnull(row.dataset):
                df = data[row.dataset]

                if pd.notnull(row.table):
                    table = getattr(package, row.table)
                    df = table(df, row.param1, row.param2)

                value = function(df, **kwargs)

            # Run methods on no datasets
            else:
                value = function(**kwargs)

            # Correct any formatting issues
            row_type = row.type.lower().strip()
            row_func = row.method.strip()

            if row_type in ['numeric', 'datetime']:
                value = calculate.remove_leading_zero(value)

                is_percent_func = row_func.startswith('percent')
                is_delta_func = row_func.startswith('delta')
                if (is_percent_func or is_delta_func) and value == 'None':
                    value = "N/A"

            # Assign value to variable
            variables.loc[index, 'value'] = calculate.format_number(value)


def build_tables(data, periods=this.build_tables.periods, progress_bar=True):
    """
    Create tables for each data element specified. The default time period
    used for this assembly is the 90-day period (d90).

    Parameters
    ----------
        data (dict): Results from `query()`.
        period (list): Keys for indicated time periods.
        progress_bar (bool): Show progress bar for this method.

    Returns
    -------
        List of DataFrames.

    Notes
    -----
        Originally, all tables were derived from the d90 dataset. Currently,
        they are derived from d90, mp0, mp1 ,mp2, and mp3 datasets. The last
        four data sets prefixed with mp- are used to get Trust scores.
        
        Columns are defined in the configuration.yaml file. Columns follow this
        format: (Name of column, limit, e.g., mininum number of respondents)
    """
    columns = this.build_tables.columns
    tables = dict()
    current = data[periods[0]].copy()
    for column, limit in columns:
        table = calculate.build_table(current, column, limit)
        table.loc[:, 'Trust'] = table.Trust.apply(calculate.add_percent)
        tables[column] = table
    return tables


def distill(text, pattern=this.distill.pattern):
    """ Get variables from a Document with the format `$NAME1`. """
    pattern = re.compile(pattern)
    return pattern.findall(text)


def distill_document(document):
    """ Get all variables in a Document's text. """
    container = set()
    for p in document.paragraphs:
        container.update(distill(p.text))

    # Updated 2019-11-18: To get variables from tables
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                container.update(distill(cell.text))
    return container


def swap(variable, variables):
    """ Return the value bound to the indicated variable name. """
    value = variables[variables.variable == variable].value
    if(value.any()):
        return value.ravel()[0]
    return None


def infuse(text, variables):
    """
    Return a string with dummy `$VARIABLES` replaced with data.

    Parameters
    ----------
        text (Paragraph): Document text.
        variables (list): Variable names to replace in Paragraph text.
    """
    if distill(text.text):
        for run in text.runs:

            # Insert A11 image (Bar Chart) and TOT image (Time Series) here
            if run.text.startswith("$IMAGEA11"):
                try:
                    value = swap(run.text, variables)
                    run.text = config.script.writer.infuse.figure_1
                    run.add_picture(value, width=Inches(7.1))
                except:
                    pass
            elif run.text.startswith("$IMAGETOT"):
                try:
                    value = swap(run.text, variables)
                    run.text = config.script.writer.infuse.figure_2
                    run.add_picture(value, width=Inches(7.1))
                except:
                    pass
            else:
                value = swap(run.text, variables)
                if(pd.notnull(value)):
                    run.text = value


def infuse_table(table, variables):
    """
    Update a current Table (non-demographic) object with data from variables.

    Parameters
    ----------
        table (Table): Document object to update with variables
        variables (DataFrame): Values to populate Table with
    """
    for row in range(1, len(table.rows)):
        for col in range(len(table.columns)):
            cell = table.cell(row, col)
            if(distill(cell.text)):
                for p in cell.paragraphs:
                    infuse(p, variables)


def update_table(table, tables):
    """
    Update a Table object with demographic data from a DataFrame.

    Parameters
    ----------
        table (Table): Document object to update with variables.
        tables (dict): DataFrame tables for each demographic.
    """
    table_name = table.cell(0, 0).text

    if table_name not in tables.keys():
        return None

    for row in range(1, len(table.rows)):

        current_table = tables[table_name]
        if(row > current_table.shape[0]):
            break

        for col in range(len(table.columns)):
            cell = table.cell(row, col)
            vals = current_table.iloc[row - 1, col]

            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = str(vals)


def compose(document, variables, tables):
    """
    Update the docx document with variables and table values.

    Parameters
    ----------
        document (Document): Object with text to update
        variables (DataFrame): Values to populate Document with
        tables (dict): DataFrame tables for each demographic

    Returns
    -------
        Document object

    Example
    -------
        >>> variables = analyze(data)
        >>> tables = build_tables(data)
        >>>
        >>> compose(document, variables, tables)
    """
    for p in document.paragraphs:
        infuse(p, variables)

    for t in document.tables:
        update_table(t, tables)
        infuse_table(t, variables)

    return document


def reset_variables(variables):
    """ Remove all assigned values from the `variables` table. """
    variables.loc[:, 'value'] = None
